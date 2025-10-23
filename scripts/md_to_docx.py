#!/usr/bin/env python3
"""Convert METHODOLOGY.md to METHODOLOGY.docx with formatting.

This script converts the methodology markdown file to a Word document
with proper formatting, maintaining headers, code blocks, tables, and
mathematical notation.
"""
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn


def setup_styles(doc):
    """Configure document styles for consistent formatting.

    Args:
        doc: python-docx Document object
    """
    # Configure Normal style
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    # Configure Heading styles
    for i in range(1, 4):
        heading_style = doc.styles[f'Heading {i}']
        heading_font = heading_style.font
        heading_font.name = 'Calibri'
        heading_font.bold = True
        heading_font.color.rgb = RGBColor(31, 73, 125)

        if i == 1:
            heading_font.size = Pt(18)
        elif i == 2:
            heading_font.size = Pt(14)
        else:
            heading_font.size = Pt(12)

    # Create or update Code style
    try:
        code_style = doc.styles['Code']
    except KeyError:
        code_style = doc.styles.add_style('Code', WD_STYLE_TYPE.PARAGRAPH)

    code_font = code_style.font
    code_font.name = 'Consolas'
    code_font.size = Pt(9)
    code_style.paragraph_format.left_indent = Inches(0.5)
    code_style.paragraph_format.space_before = Pt(6)
    code_style.paragraph_format.space_after = Pt(6)


def parse_markdown_line(line):
    """Parse a markdown line and determine its type.

    Args:
        line: String containing a line of markdown

    Returns:
        Tuple of (line_type, content, level) where:
        - line_type: 'heading', 'code', 'bullet', 'numbered', 'table', 'text', 'blank'
        - content: Processed text content
        - level: Heading level (1-6) or bullet depth (0-n)
    """
    line = line.rstrip()

    if not line:
        return ('blank', '', 0)

    # Headings
    heading_match = re.match(r'^(#{1,6})\s+(.+)$', line)
    if heading_match:
        level = len(heading_match.group(1))
        content = heading_match.group(2)
        return ('heading', content, level)

    # Code block markers
    if line.strip().startswith('```'):
        return ('code_fence', line.strip()[3:], 0)

    # Bullets
    bullet_match = re.match(r'^(\s*)([-*+])\s+(.+)$', line)
    if bullet_match:
        indent = len(bullet_match.group(1))
        content = bullet_match.group(3)
        level = indent // 2
        return ('bullet', content, level)

    # Numbered lists
    numbered_match = re.match(r'^(\s*)(\d+\.)\s+(.+)$', line)
    if numbered_match:
        indent = len(numbered_match.group(1))
        content = numbered_match.group(3)
        level = indent // 2
        return ('numbered', content, level)

    # Table rows
    if '|' in line and not line.strip().startswith('|'):
        # Likely a table row
        return ('table', line, 0)

    # Table separator (ignore)
    if re.match(r'^\s*\|[-:\s|]+\|\s*$', line):
        return ('table_sep', '', 0)

    # Horizontal rule
    if re.match(r'^[-*_]{3,}$', line.strip()):
        return ('hr', '', 0)

    # Regular text
    return ('text', line, 0)


def process_inline_formatting(text):
    """Process inline markdown formatting (bold, italic, code).

    Args:
        text: String with markdown formatting

    Returns:
        List of (text, format_dict) tuples where format_dict contains
        'bold', 'italic', 'code' boolean flags
    """
    # This is a simplified parser - for production, use a proper markdown parser
    # For now, return plain text segments
    segments = []

    # Handle bold (**text** or __text__)
    parts = re.split(r'(\*\*.*?\*\*|__.*?__)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            segments.append((part[2:-2], {'bold': True, 'italic': False, 'code': False}))
        elif part.startswith('__') and part.endswith('__'):
            segments.append((part[2:-2], {'bold': True, 'italic': False, 'code': False}))
        elif part:
            # Check for italic or code
            if part.startswith('`') and part.endswith('`'):
                segments.append((part[1:-1], {'bold': False, 'italic': False, 'code': True}))
            elif part.startswith('*') and part.endswith('*') and not part.startswith('**'):
                segments.append((part[1:-1], {'bold': False, 'italic': True, 'code': False}))
            else:
                segments.append((part, {'bold': False, 'italic': False, 'code': False}))

    return segments if segments else [(text, {'bold': False, 'italic': False, 'code': False})]


def add_formatted_paragraph(doc, text, style_name='Normal'):
    """Add a paragraph with inline formatting to the document.

    Args:
        doc: python-docx Document
        text: Text with markdown inline formatting
        style_name: Style to apply to the paragraph

    Returns:
        The created paragraph object
    """
    para = doc.add_paragraph(style=style_name)
    segments = process_inline_formatting(text)

    for segment_text, formatting in segments:
        run = para.add_run(segment_text)
        if formatting['bold']:
            run.bold = True
        if formatting['italic']:
            run.italic = True
        if formatting['code']:
            run.font.name = 'Consolas'
            run.font.size = Pt(9)

    return para


def convert_md_to_docx(md_path, docx_path):
    """Convert markdown file to formatted Word document.

    Args:
        md_path: Path to input markdown file
        docx_path: Path to output docx file
    """
    # Read markdown content
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    # Create document
    doc = Document()
    setup_styles(doc)

    # Set document properties
    doc.core_properties.title = 'Survival Analysis Methodology'
    doc.core_properties.author = 'Survival Analysis Framework'

    # Parse and convert
    in_code_block = False
    code_buffer = []
    table_buffer = []
    in_table = False

    for line in lines:
        line_type, content, level = parse_markdown_line(line)

        # Handle code blocks
        if line_type == 'code_fence':
            if in_code_block:
                # End of code block - flush buffer
                if code_buffer:
                    code_text = '\n'.join(code_buffer)
                    para = doc.add_paragraph(code_text, style='Code')
                    # Add gray background
                    shading_elm = parse_xml(r'<w:shd {} w:fill="F0F0F0"/>'.format(
                        nsdecls('w')))
                    para._element.get_or_add_pPr().append(shading_elm)
                code_buffer = []
                in_code_block = False
            else:
                # Start of code block
                in_code_block = True
            continue

        if in_code_block:
            code_buffer.append(line.rstrip())
            continue

        # Handle tables
        if line_type == 'table':
            if not in_table:
                in_table = True
                table_buffer = []
            table_buffer.append(content)
            continue
        elif in_table and line_type != 'table_sep':
            # End of table - process buffer
            if table_buffer:
                process_table(doc, table_buffer)
            table_buffer = []
            in_table = False

        if line_type == 'table_sep':
            continue

        # Process non-table content
        if line_type == 'heading':
            if level <= 3:
                doc.add_heading(content, level=level)
            else:
                # For h4-h6, use bold text
                para = doc.add_paragraph()
                run = para.add_run(content)
                run.bold = True
                run.font.size = Pt(11)

        elif line_type == 'bullet':
            add_formatted_paragraph(doc, content, 'List Bullet')

        elif line_type == 'numbered':
            add_formatted_paragraph(doc, content, 'List Number')

        elif line_type == 'hr':
            doc.add_paragraph('_' * 80)

        elif line_type == 'text':
            add_formatted_paragraph(doc, content)

        elif line_type == 'blank':
            # Only add blank line if previous wasn't blank
            if doc.paragraphs and doc.paragraphs[-1].text.strip():
                doc.add_paragraph()

    # Save document
    doc.save(docx_path)
    print(f"Converted {md_path} -> {docx_path}")


def process_table(doc, table_lines):
    """Process table lines and add table to document.

    Args:
        doc: python-docx Document
        table_lines: List of table row strings
    """
    if not table_lines:
        return

    # Parse table structure
    rows = []
    for line in table_lines:
        cells = [cell.strip() for cell in line.split('|')]
        # Remove empty first/last cells from split
        cells = [c for c in cells if c]
        if cells:
            rows.append(cells)

    if not rows:
        return

    # Create table
    num_cols = len(rows[0])
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.style = 'Light Grid Accent 1'

    # Populate table
    for i, row_data in enumerate(rows):
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = table.rows[i].cells[j]
                cell.text = cell_text
                # Bold first row (header)
                if i == 0:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True


def nsdecls(*prefixes):
    """Generate namespace declarations for XML."""
    return ' '.join(f'xmlns:{p}="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
                    for p in prefixes)


def parse_xml(xml_string):
    """Parse XML string into an element."""
    from lxml import etree
    return etree.fromstring(xml_string)


if __name__ == '__main__':
    # Get paths relative to script location
    script_dir = Path(__file__).parent
    repo_root = script_dir.parent

    md_path = repo_root / 'METHODOLOGY.md'
    docx_path = repo_root / 'METHODOLOGY.docx'

    if not md_path.exists():
        print(f"Error: {md_path} not found")
        exit(1)

    convert_md_to_docx(str(md_path), str(docx_path))
    print(f"Successfully created {docx_path}")
